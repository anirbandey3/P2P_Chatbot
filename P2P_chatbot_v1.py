import streamlit as st
import pandas as pd
import google.generativeai as genai
import matplotlib.pyplot as plt
import io
import os
import json
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import requests
from googleapiclient.discovery import build
import requests
from bs4 import BeautifulSoup
import time
import json
from urllib.parse import urljoin
import re


# Configure page settings
st.set_page_config(page_title="P2P Chatbot", layout="wide")

GOOGLE_API_KEY = "AIzaSyC79GAiVjR-hFbIGxne7Fp0duT8bs8wniE"
GOOGLE_CSE_ID = "22b40f99ada7e4fc1" 
GEMINI_API_KEY = "AIzaSyDbpQo8dISvrlrph3EGf-CDGgoBN-EBCxo"
genai.configure(api_key=GEMINI_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-1.5-pro')

# Initialize session state variables
if "messages" not in st.session_state:
    st.session_state["messages"] = []
    
if "feedback_data" not in st.session_state:
    st.session_state["feedback_data"] = {}

if "current_query" not in st.session_state:
    st.session_state["current_query"] = None
    
if "awaiting_regeneration" not in st.session_state:
    st.session_state["awaiting_regeneration"] = False

if "query_attempts" not in st.session_state:
    st.session_state["query_attempts"] = {}

if "doc_data" not in st.session_state:
    st.session_state["doc_data"] = None

if "vectorstore" not in st.session_state:
    st.session_state["vectorstore"] = None

# File to store feedback data
FEEDBACK_FILE = "feedback_data.json"

def load_feedback_data():
    """Load feedback data from file if it exists"""
    try:
        if os.path.exists(FEEDBACK_FILE):
            with open(FEEDBACK_FILE, 'r') as f:
                return json.load(f)
    except Exception as e:
        st.error(f"Error loading feedback data: {e}")
    return {}

def save_feedback_data():
    """Save feedback data to file"""
    try:
        with open(FEEDBACK_FILE, 'w') as f:
            json.dump(st.session_state.feedback_data, f)
    except Exception as e:
        st.error(f"Error saving feedback data: {e}")

def rate_response(message_id, star_num):
    """Handle star rating logic"""
    # Get the current query
    current_query = st.session_state.current_query
    
    # Save the rating
    if message_id not in st.session_state.feedback_data:
        st.session_state.feedback_data[message_id] = {}
        
    # Get the message data
    message = st.session_state.messages[message_id]
    
    # Store feedback data
    st.session_state.feedback_data[message_id].update({
        "query": current_query,
        "code": message.get("code", ""),
        "results": str(message.get("results", "")),
        "rating": star_num,
        "context": get_recent_conversation_context()
    })
    
    # Set flag for regeneration if rating is less than 3
    if star_num < 3:
        # Increment attempt counter
        if current_query not in st.session_state.query_attempts:
            st.session_state.query_attempts[current_query] = 1
        else:
            st.session_state.query_attempts[current_query] += 1
            
        # Set regeneration flag
        st.session_state.awaiting_regeneration = True
    
    # Save to file
    save_feedback_data()

def get_recent_conversation_context():
    """Get the recent conversation context"""
    context = []
    recent_messages = st.session_state.messages[-5:] if len(st.session_state.messages) > 5 else st.session_state.messages
    
    for msg in recent_messages:
        if msg["role"] == "user":
            context.append(f"User: {msg['content']}")
        else:
            context.append(f"Assistant: {msg['content']}")
    
    return "\n".join(context)

def display_chat_message(role, content, code=None, results=None, recommendations=None, message_id=None):
    """Display a chat message with optional code, results, recommendations, and star rating."""
    with st.chat_message(role):
        st.markdown(content)
        
        if code is not None:
            st.markdown("**Generated Python Code:**")
            st.code(code, language="python")
            
            # Add star rating if this is an assistant message
            if role == "assistant" and message_id is not None:
                # Check if feedback was already given for this message
                if message_id in st.session_state.feedback_data and "rating" in st.session_state.feedback_data[message_id]:
                    rating = st.session_state.feedback_data[message_id]["rating"]
                    st.success(f"You rated this response: {rating}/5 stars")
                else:
                    st.markdown("**How helpful was this response?**")
                    cols = st.columns(5)
                    
                    for i in range(5):
                        star_num = i + 1
                        star_key = f"star_{star_num}_{message_id}"
                        
                        with cols[i]:
                            if st.button(f"⭐ {star_num}", key=star_key):
                                rate_response(message_id, star_num)
                                st.rerun()
        
        if results is not None:
            st.markdown("**Results:**")
            st.markdown(str(results))

@st.cache_resource
def get_embeddings_model():
    """Initialize and cache the embedding model"""
    try:
        return HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        # Test the model
        test_embedding = model.embed_query("test")
        if test_embedding is None:
            raise ValueError("Embeddings model failed to generate embeddings")
        return model
    except Exception as e:
        st.error(f"Failed to initialize embeddings model: {str(e)}")
        return None
# Add this function to detect schema from the user's question
def detect_schema_from_query(query):
    """Detect schema name from user query.
    Returns the detected schema name or 'AP' as default."""
    
    # Extract common Oracle schemas from the JSON file
    schemas = []
    try:
        with open("schema_table_links.json", "r") as f:
            schema_data = json.load(f)
            schemas = list(schema_data.keys())
    except Exception as e:
        print(f"Error reading schemas from JSON: {e}")
        # Fallback to common schemas if file can't be read
        schemas = ["AP", "GL", "PO", "AR", "FA", "INV", "CE", "HR", "FND", "AHL", "AK", "ALR", "AMS"]
    
    # First look for direct schema mentions in the query
    # Look for exact schema names (prioritize exact matches)
    for schema in schemas:
        pattern = fr'\b{schema}\b'
        if re.search(pattern, query, re.IGNORECASE):
            return schema.upper()
    
    # Check common patterns that might indicate a schema
    schema_patterns = {
        r'\b(?:in|from|using|for)\s+(\w+)\s+schema\b': 1,  # "in XX schema"
        r'\bschema\s+(\w+)\b': 1,                          # "schema XX"
        r'\b(\w+)\.[\w_]+\b': 1,                           # "XX.TABLE_NAME"
    }
    
    # Check each pattern
    for pattern, group in schema_patterns.items():
        matches = re.finditer(pattern, query, re.IGNORECASE)
        for match in matches:
            schema = match.group(group).upper()
            # Validate if this matched schema exists in our list
            if schema in [s.upper() for s in schemas]:
                return schema
    
    # Special case handling for common patterns
    if "GL_" in query.upper() or "general ledger" in query.lower():
        return "GL"
    elif "AP_" in query.upper() or "payables" in query.lower():
        return "AP"
    elif "AR_" in query.upper() or "receivables" in query.lower():
        return "AR"
    elif "PO_" in query.upper() or "purchase order" in query.lower():
        return "PO"
    
    # Default to AP if no schema detected
    return "AP"



def fetch_table_columns_from_json(table_name, schema_name="AP", json_file="schema_table_links.json"):
    """Fetch column details for a specific table using a JSON file with schema and table URLs."""
    try:
        # Load the JSON file
        with open(json_file, "r") as f:
            schema_data = json.load(f)

        # Check if the schema exists in the JSON file
        if schema_name not in schema_data:
            return f"No schema found for '{schema_name}'. Please verify the schema name."

        # Find the URL for the requested table
        table_url = None
        for url in schema_data[schema_name]:
            if f"c_name={table_name.upper()}" in url:
                table_url = url
                break

        if not table_url:
            return f"No table found for '{table_name}' in schema '{schema_name}'. Please verify the table name."

        # Fetch the table details page
        print(f"Fetching table details from URL: {table_url}")  # Debugging
        response = requests.get(table_url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        # Find the table containing column details
        column_table = soup.find("table", {"summary": "Column details for this table"})
        if not column_table:
            return f"No column details found for table '{table_name}'."

        # Extract column details
        columns = []
        rows = column_table.find_all("tr")[1:]  # Skip the header row
        for row in rows:
            cells = row.find_all("td")
            if len(cells) >= 5:
                column_name = cells[0].text.strip()
                data_type = cells[1].text.strip()
                length = cells[2].text.strip()
                nullable = cells[3].text.strip()
                description = cells[4].text.strip()
                columns.append({
                    "name": column_name,
                    "type": data_type,
                    "length": length,
                    "nullable": nullable,
                    "description": description
                })

        # Format the response
        if columns:
            formatted_columns = "\n".join(
                [f"- **{col['name']}** ({col['type']}): Length={col['length']}, Nullable={col['nullable']}, Description={col['description']}" for col in columns]
            )
            return f"### Columns in {table_name}\n\n{formatted_columns}"
        else:
            return f"No column details found for table '{table_name}'."
    except FileNotFoundError:
        return f"JSON file '{json_file}' not found. Please ensure the file exists."
    except Exception as e:
        return f"Failed to fetch table columns: {str(e)}"

def get_relevant_chunks(query, doc_data):
    """Retrieve relevant text chunks using embeddings and similarity search."""
    if not doc_data or 'chunks' not in doc_data or not doc_data['chunks']:
        st.error("No document content available")
        return []

    try:
        # Get cached embedding model
        embeddings_model = get_embeddings_model()
        if embeddings_model is None:
            st.error("Failed to initialize embeddings model")
            return []

        # Create or get vectorstore
        if "vectorstore" not in st.session_state or st.session_state.vectorstore is None:
            with st.spinner("Creating embeddings..."):
                try:
                    # Create ChromaDB directory if it doesn't exist
                    os.makedirs("./chroma_db", exist_ok=True)
                    
                    # Create new vectorstore with persist
                    st.session_state.vectorstore = Chroma(
                        persist_directory="./chroma_db",
                        embedding_function=embeddings_model
                    )
                    
                    # Add documents to vectorstore
                    st.session_state.vectorstore.add_texts(
                        texts=doc_data['chunks']
                    )
                    
                    st.success("Embeddings created successfully!")
                except Exception as e:
                    st.error(f"Failed to create vectorstore: {str(e)}")
                    return []

        # Perform similarity search
        try:
            results = st.session_state.vectorstore.similarity_search(
                query,
                k=5  # Reduced from 5 to 3 for better relevance
            )
            if not results:
                st.warning("No relevant content found")
                return []
            
            return [doc.page_content for doc in results]
        except Exception as e:
            st.error(f"Search failed: {str(e)}")
            return []

    except Exception as e:
        st.error(f"Error processing embeddings: {str(e)}")
        return []


def process_doc():
    """Process the DOC file and store its content"""
    try:
        file_path = os.path.join(os.path.dirname(__file__), "Payables User Guide.docx")
        if os.path.exists(file_path):
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len
            )
            chunks = text_splitter.split_text(text)
            
            st.session_state.doc_data = {
                'text': text,
                'chunks': chunks,
                'filename': os.path.basename(file_path),
                'summary': f"Document contains {len(doc.paragraphs)} paragraphs and {len(chunks)} text chunks."
            }
            return True
    except Exception as e:
        st.error(f"Error processing DOC file: {e}")
        return False
    
def get_web_content(query, num_results=3):
    """Get relevant content using Google Custom Search API"""
    web_content = []
    
    try:
        # Initialize the Custom Search API service
        service = build(
            "customsearch", "v1",
            developerKey=GOOGLE_API_KEY
        )
        
        # Perform the search
        result = service.cse().list(
            q=query,
            cx=GOOGLE_CSE_ID,
            num=num_results
        ).execute()
        
        # Process search results
        if 'items' in result:
            for item in result['items']:
                web_content.append({
                    'url': item['link'],
                    'title': item.get('title', ''),
                    'content': item.get('snippet', ''),
                    'source': 'Google Custom Search'
                })
                
        return web_content
    except Exception as e:
        st.error(f"Google Search API failed: {str(e)}")
        return []

def main():
    st.title("P2P Q&A App (Payables User Guide)")
    
    # Load feedback data
    if not st.session_state["feedback_data"]:
        st.session_state["feedback_data"] = load_feedback_data()
    
    # Process DOC file
    if st.session_state["doc_data"] is None:
        if process_doc():
            st.success("Document processed successfully!")
            st.info(st.session_state["doc_data"]["summary"])
    
    if st.session_state["doc_data"] is not None:
        
        # Sidebar with feedback history
        if st.sidebar.checkbox("Show Feedback History"):
            st.sidebar.markdown("### Feedback History")
            for msg_id, data in st.session_state.feedback_data.items():
                if "rating" in data and data.get("rating", 0) >= 3:
                    st.sidebar.markdown(f"**Query:** {data['query']}")
                    st.sidebar.markdown(f"**Rating:** {data['rating']}/5 stars")
                    st.sidebar.code(data['code'], language="python")
                    st.sidebar.markdown(f"**Context:** {data.get('context', 'No context available')}")
                    st.sidebar.divider()

        # Display chat history
        for i, message in enumerate(st.session_state.messages):
            display_chat_message(
                message["role"],
                message["content"],
                message.get("code"),
                message.get("results"),
                i if message["role"] == "assistant" else None
            )

        # Check if we need to regenerate
        if st.session_state.awaiting_regeneration:
            with st.spinner("Generating an improved response based on your feedback..."):
                # Get the current attempt number
                query = st.session_state.current_query
                attempt_num = st.session_state.query_attempts.get(query, 1)

        # Chat input
        if query := st.chat_input("Ask me about Payables..."):
            st.session_state.current_query = query
            display_chat_message("user", query)
            st.session_state.messages.append({"role": "user", "content": query})
            
            with st.spinner("Analyzing documentation and web sources..."):
                # Check if the query is about a specific table
                schema_table_pattern = r'\b([A-Z_]+)\.([A-Z_]+_ALL|[A-Z_]+_INTERFACE|[A-Z_]+_HEADERS|[A-Z_]+_LINES)\b'
                schema_table_match = re.search(schema_table_pattern, query.upper())
                #schema_pattern = r'\b(AD_MONITOR|AHL|AK|ALR|AMS|AMV|ANONYMOUS|AP|APPLSYS|APPLSYSPUB|APPQOSSYS|APPS|APPS_NE|AR|ASF|ASG|ASL|ASN|ASO|ASP|AST|AUDSYS|AX|AZ|BEN|BIC|BIM|BIS|BNE|BOM|CCT|CE|CLN|CN|CRP|CS|CSC|CSD|CSE|CSF|CSI|CSL|CSM|CSP|CSR|CTXSYS|CUA|CUG|CZ|DBSNMP|DBVSECOPS|DBV_ACCTMGR|DBV_OWNER|DDR|DIP|DNA|DOM|DPP|DP_MGR|DVF|DVSYS|EAM|EC|ECX|EDR|EGO|EM_MONITOR|ENG|ENI|ETRM|ETRM_READONLY|FA|FLM|FPA|FRM|FTE|FUN|FV|GHG|GL|GMA|GMD|GME|GMF|GMI|GML|GMO|GMP|GMS|GR|GSMADMIN_INTERNAL|GSMCATUSER|GSMUSER|HR|HRI|HXC|HXT|IA|IBC|IBE|IBU|IBW|IBY|ICX|IEB|IEC|IEM|IEO|IES|IEU|IEX|IGC|IGI|IMC|INL|INV|IPA|IPM|ITG|IZU|JA|JE|JG|JL|JMF|JTF|JTM|LBACSYS|LNS|MDDATA|MDSYS|MFG|MGDSYS|MRP|MSC|MSD|MSO|MSR|MTH|MWA|ODM|ODM_MTR|OE|OJVMSYS|OKC|OKE|OKL|OKS|OKX|ONT|OPI|ORACLE_OCM|ORDDATA|ORDPLUGINS|ORDSYS|OSM|OTA|OUTLN|OZF|PA|PJI|PJM|PMI|PN|PO|POM|PON|POS|PRP|PSA|PSP|PV|QA|QOT|QP|QPR|QRM|RG|RLM|RRS|SCOTT|SECADMIN|SI_INFORMTN_SCHEMA|SPATIAL_CSW_ADMIN_USR|SPATIAL_WFS_ADMIN_USR|SSOSDK|SSP|SYS|SYSBACKUP|SYSDG|SYSKM|SYSTEM|VEA|WIP|WMS|WPS|WSH|WSM|XDB|XDO|XDP|XLA|XLE|XNB|XNP|XS\$NULL|XTR|ZX)\b'
                #table_match = re.search(schema_table_pattern, query.upper())
                
      
                if schema_table_match:
                    schema = schema_table_match.group(1)
                    table_name = schema_table_match.group(2)
                    # schema_name = detect_schema_from_query(query)
                    response = fetch_table_columns_from_json(table_name,schema_name=schema)
                else:
                    # Get doc content
                    doc_chunks = get_relevant_chunks(query, st.session_state.doc_data)
                    
                    # Get web content
                    web_content = get_web_content(query)
                    
                    # Generate enhanced response
                    response = generate_enhanced_response(query, doc_chunks, web_content)
                
                if response:
                    message_id = len(st.session_state.messages)
                    st.session_state.messages.append({
                        "role": "assistant",
                        "content": response,
                        "results": None,
                        "recommendations": None
                    })
                    display_chat_message("assistant", response)

        # Sidebar buttons
        if st.sidebar.button("Clear Chat History"):
            st.session_state.messages = []
            st.session_state.current_query = None
            st.session_state.awaiting_regeneration = False
            st.session_state.query_attempts = {}
            st.rerun()
            
        if st.sidebar.button("Clear Feedback Data"):
            st.session_state.feedback_data = {}
            if os.path.exists(FEEDBACK_FILE):
                os.remove(FEEDBACK_FILE)
            st.rerun()

def generate_enhanced_response(query, doc_chunks, web_content, table_name=None):
    """Generate response for DOC queries using Gemini model."""
    if not st.session_state.doc_data:
        return "No Document data available. Please load a Doc file first."

    # Include table_name in the prompt only if it is provided
    table_section = f"""
## Table: {table_name}
| Column Name   | Data Type   | Length | Nullable | Description          |
|---------------|-------------|--------|----------|----------------------|
| INVOICE_ID    | NUMBER      | (15)   | Yes      | Invoice identifier   |
| VENDOR_ID     | NUMBER      | (10)   | No       | Vendor identifier    |
""" if table_name else ""

    prompt = f"""
You are an expert AI assistant specializing in Oracle Payables and P2P processes. When the user asks about a specific table, fetch the column details from the Oracle ETRM website and provide a detailed response.

For table queries:
1. Provide the column names, data types, lengths, nullable status, and descriptions.
2. Format the response as a markdown table for clarity.

For other queries:
1. Combine information from the official documentation and web sources.
2. Provide a comprehensive answer with clear explanations, examples, and best practices.

User Question: {query}

{table_section}

If the query is not about a table, follow this structure:

Document being analyzed: {st.session_state.doc_data.get('filename', 'Unknown')}

Document Content:
{doc_chunks}

Web Sources:
{[f"Source: {content['url']}\n{content['content']}" for content in web_content]}

Please provide a detailed response following this structure:

## Main Answer
- Provide a thorough, detailed explanation.
- Break down complex concepts into clear points.
- Include specific details and examples from the document.
- Use bullet points and numbered lists for clarity.

## More Knowledge
- Include insights from web sources.
- Reference the source URLs.
- Highlight industry best practices.
- Add real-world examples and use cases.

## Supporting Evidence
- Quote relevant passages from the document (use > for quotes).
- Explain how each quote supports your answer.
- Reference specific sections or pages when available.

## Additional Context
- Provide background information if relevant.
- Explain related concepts that help understand the topic.
- Make connections to other relevant topics in the document.

## Practical Applications
- Include real-world examples or use cases.
- Explain how this information is applied in practice.
- Mention best practices or common scenarios.

## Follow-up Questions
- Suggest 3-4 detailed follow-up questions.
- Focus on questions that would deepen understanding.
- Include questions about related topics.

Format your response using clear markdown formatting:
- Use ## for section headers.
- Use * or - for bullet points.
- Use > for quotes.
- Use **bold** for emphasis.
- Use numbered lists for sequential information.

If information conflicts between sources, explain the differences and recommend the best approach.
"""

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating response: {e}")
        return None

if __name__ == "__main__":
    main()